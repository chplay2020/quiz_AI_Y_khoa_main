"""
Question Generation API Routes
Handles question generation, retrieval, and management
"""
from unittest import result
from fastapi import APIRouter, HTTPException, BackgroundTasks
from typing import Optional, List
import uuid
from datetime import datetime
import asyncio
import structlog
import json
import os
from pathlib import Path

from app.config import settings
from app.models import (
    QuestionGenerationRequest, QuestionGenerationResponse,
    GenerationStatus, QuestionResponse, APIResponse,
    SearchRequest, SearchResponse, SearchResult,
    QuestionStats, ExportRequest, ExportResponse, ExportFormat, ExportMode
)
from app.core.rag_engine import get_rag_engine
from app.core.llm_provider import get_llm_provider
from app.flows import create_question_generation_flow

logger = structlog.get_logger()
router = APIRouter(prefix="/questions", tags=["Questions"])

# Persistent storage paths
QUESTIONS_DB_FILE = Path(settings.DATA_DIR) / "questions_db.json"
TASKS_DB_FILE = Path(settings.DATA_DIR) / "generation_tasks.json"

# In-memory storage (with persistent backup)
questions_db = {}
generation_tasks = {}


def load_persistent_data():
    """Load data from persistent storage"""
    global questions_db, generation_tasks
    
    # Load questions
    if QUESTIONS_DB_FILE.exists():
        try:
            with open(QUESTIONS_DB_FILE, 'r', encoding='utf-8') as f:
                questions_db = json.load(f)
            logger.info("Loaded questions from persistent storage", count=len(questions_db))
        except Exception as e:
            logger.error("Failed to load questions", error=str(e))
    
    # Load tasks
    if TASKS_DB_FILE.exists():
        try:
            with open(TASKS_DB_FILE, 'r', encoding='utf-8') as f:
                generation_tasks = json.load(f)
            logger.info("Loaded tasks from persistent storage", count=len(generation_tasks))
        except Exception as e:
            logger.error("Failed to load tasks", error=str(e))


def save_questions_db():
    """Save questions to persistent storage"""
    try:
        QUESTIONS_DB_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(QUESTIONS_DB_FILE, 'w', encoding='utf-8') as f:
            json.dump(questions_db, f, ensure_ascii=False, indent=2)
        logger.debug("Saved questions to persistent storage", count=len(questions_db))
    except Exception as e:
        logger.error("Failed to save questions", error=str(e))


def save_tasks_db():
    """Save tasks to persistent storage"""
    try:
        TASKS_DB_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(TASKS_DB_FILE, 'w', encoding='utf-8') as f:
            json.dump(generation_tasks, f, ensure_ascii=False, indent=2)
        logger.debug("Saved tasks to persistent storage", count=len(generation_tasks))
    except Exception as e:
        logger.error("Failed to save tasks", error=str(e))


# Load data on startup
load_persistent_data()


async def generate_questions_background(
    task_id: str,
    request: QuestionGenerationRequest
):
    """Background task for question generation"""
    try:
        generation_tasks[task_id]['status'] = 'processing'
        
        # Get LLM provider
        llm = get_llm_provider()
        
        # Create and run the question generation flow
        flow = create_question_generation_flow(
            llm_provider=llm,
            include_case_based=request.include_case_based,
            enable_double_check=False,  # AI double-check disabled
            use_gemini_thinking=getattr(request, 'use_gemini_thinking', False)
        )
        
        logger.info("Starting flow execution", task_id=task_id, 
                   use_gemini_thinking=getattr(request, 'use_gemini_thinking', False))

        result = await flow.run({
            'document_ids': request.document_ids,
            'num_questions': request.num_questions,
            'difficulty': request.difficulty.value if request.difficulty else 'medium',
            'question_types': [qt.value for qt in request.question_types] if request.question_types else ['single_choice'],
            'topics': request.topics,
            'focus_areas': request.focus_areas,
            'language': 'vi',  # Luôn sử dụng tiếng Việt
            'include_case_based': request.include_case_based,  # Truyền vào để tính số lượng đúng
            'use_google_search': getattr(request, 'use_google_search', False)
        })
        
        logger.info("Flow execution completed", task_id=task_id)

        
        # ✅ FIX: Check for flow failure status (fail-fast from NoContextErrorNode)
        if result.get("status") == "failed":
            error_msg = result.get("error", "Không tìm thấy nội dung hợp lệ trong tài liệu")
            logger.error("Flow failed - no valid context", task_id=task_id, error=error_msg)
            generation_tasks[task_id]["status"] = "failed"
            generation_tasks[task_id]["questions"] = []
            generation_tasks[task_id]["total_questions"] = 0
            generation_tasks[task_id]["generated_questions"] = 0
            generation_tasks[task_id]["progress"] = 1.0
            generation_tasks[task_id]["error"] = error_msg
            return

        # ✅ FIX TREO TASK KHI KHÔNG CÓ CONTEXT
        if not result.get("retrieved_contexts"):
            logger.error("NO CONTEXT RETRIEVED – END TASK EARLY", task_id=task_id)
            generation_tasks[task_id]["status"] = "failed"
            generation_tasks[task_id]["questions"] = []
            generation_tasks[task_id]["total_questions"] = 0
            generation_tasks[task_id]["generated_questions"] = 0
            generation_tasks[task_id]["progress"] = 1.0
            generation_tasks[task_id]["error"] = "Không tìm thấy nội dung y khoa hợp lệ. Tài liệu có thể chỉ chứa watermark hoặc là file scan chất lượng thấp."
            return
        
        # Debug log
        logger.info("Flow result received", result_type=type(result).__name__, 
                   has_reviewed=('reviewed_questions' in result) if isinstance(result, dict) else False,
                   has_validated=('validated_questions' in result) if isinstance(result, dict) else False,
                   result_keys=list(result.keys()) if isinstance(result, dict) else "not_dict")
        
        # Ensure result is a dict
        if not isinstance(result, dict):
            logger.error("Flow returned non-dict result", result_type=type(result).__name__, result_content=str(result)[:500])
            raise ValueError(f"Invalid flow result type: {type(result).__name__}")
        
        # Get reviewed questions (with AI double-check results)
        questions = result.get('reviewed_questions', result.get('validated_questions', []))
        # 🔍 DEBUG: kiểm tra ngôn ngữ câu hỏi thực tế
        if questions:
            q = questions[0]
            logger.warning(
                "SAMPLE QUESTION",
                question_text=q.get("question_text", "")[:200]
            )
        review_stats = result.get('review_stats', {})
        
        # Store questions
        stored_questions = []
        for q in questions:
            question_id = str(uuid.uuid4())
            q['id'] = question_id
            q['created_at'] = datetime.utcnow().isoformat()
            questions_db[question_id] = q
            stored_questions.append(q)
        
        logger.info(
            "Stored questions in database",
            task_id=task_id,
            num_stored=len(stored_questions),
            question_ids=[q['id'] for q in stored_questions],
            total_in_db=len(questions_db)
        )
        
        # Save to persistent storage
        save_questions_db()
        
        # Update task status
        generation_tasks[task_id]['status'] = 'completed'
        generation_tasks[task_id]['questions'] = stored_questions
        generation_tasks[task_id]['total_questions'] = len(stored_questions)
        generation_tasks[task_id]['generated_questions'] = len(stored_questions)
        generation_tasks[task_id]['progress'] = 1.0
        generation_tasks[task_id]['review_stats'] = review_stats
        
        # Save tasks to persistent storage
        save_tasks_db()
        
        logger.info(
            "Question generation completed",
            task_id=task_id,
            num_questions=len(stored_questions),
            review_stats=review_stats
        )
        
    except Exception as e:
        import traceback
        logger.error("Question generation failed", 
                    task_id=task_id, 
                    error=str(e),
                    error_type=type(e).__name__,
                    traceback=traceback.format_exc())
        generation_tasks[task_id]['status'] = 'failed'
        generation_tasks[task_id]['error'] = str(e)


@router.post("/generate", response_model=QuestionGenerationResponse)
async def generate_questions(
    background_tasks: BackgroundTasks,
    request: QuestionGenerationRequest
):
    """
    Start question generation from documents.
    Returns a task ID that can be used to check progress.
    """
    # Validate request
    if not request.document_ids:
        raise HTTPException(status_code=400, detail="At least one document ID is required")
    
    if request.num_questions > settings.MAX_QUESTIONS_PER_REQUEST:
        raise HTTPException(
            status_code=400,
            detail=f"Maximum {settings.MAX_QUESTIONS_PER_REQUEST} questions per request"
        )
    
    # Create task
    task_id = str(uuid.uuid4())
    generation_tasks[task_id] = {
        'task_id': task_id,
        'status': 'pending',
        'progress': 0.0,
        'total_questions': request.num_questions,
        'generated_questions': 0,
        'questions': None,
        'error': None,
        'created_at': datetime.utcnow().isoformat()
    }
    
    # Save tasks
    save_tasks_db()
    
    # Start background task
    background_tasks.add_task(
        generate_questions_background,
        task_id,
        request
    )
    
    return QuestionGenerationResponse(
        task_id=task_id,
        status='pending',
        message='Question generation started'
    )


@router.get("/generate/{task_id}/status", response_model=GenerationStatus)
async def get_generation_status(task_id: str):
    """Get the status of a question generation task"""
    if task_id not in generation_tasks:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task = generation_tasks[task_id]
    
    return GenerationStatus(
        task_id=task_id,
        status=task['status'],
        progress=task['progress'],
        total_questions=task['total_questions'],
        generated_questions=task['generated_questions'],
        questions=task.get('questions'),
        error=task.get('error')
    )


@router.get("/", response_model=APIResponse)
async def list_questions(
    document_id: Optional[str] = None,
    difficulty: Optional[str] = None,
    question_type: Optional[str] = None,
    topic: Optional[str] = None,
    limit: int = 50,
    offset: int = 0
):
    """List all generated questions with optional filtering"""
    questions = list(questions_db.values())
    
    # Apply filters
    if document_id:
        questions = [q for q in questions if q.get('document_id') == document_id]
    if difficulty:
        questions = [q for q in questions if q.get('difficulty') == difficulty]
    if question_type:
        questions = [q for q in questions if q.get('question_type') == question_type]
    if topic:
        questions = [q for q in questions if topic.lower() in q.get('topic', '').lower()]
    
    # Pagination
    total = len(questions)
    questions = questions[offset:offset + limit]
    
    return APIResponse(
        success=True,
        data={
            'questions': questions,
            'total': total,
            'limit': limit,
            'offset': offset
        }
    )


@router.get("/{question_id}", response_model=APIResponse)
async def get_question(question_id: str):
    """Get a specific question by ID"""
    if question_id not in questions_db:
        raise HTTPException(status_code=404, detail="Question not found")
    
    return APIResponse(
        success=True,
        data=questions_db[question_id]
    )


@router.put("/{question_id}", response_model=APIResponse)
async def update_question(question_id: str, updates: dict):
    """Update a question"""
    logger.info(
        "Update question request",
        question_id=question_id,
        updates_keys=list(updates.keys()),
        total_questions_in_db=len(questions_db),
        question_exists=question_id in questions_db
    )
    
    if question_id not in questions_db:
        # Log all available IDs for debugging
        logger.error(
            "Question not found in database",
            question_id=question_id,
            available_ids=list(questions_db.keys())[:10]  # First 10 IDs
        )
        raise HTTPException(status_code=404, detail="Question not found")
    
    question = questions_db[question_id]
    
    # Update allowed fields
    allowed_fields = ['question_text', 'options', 'correct_answer', 'explanation', 
                      'difficulty', 'topic', 'keywords']
    
    for field in allowed_fields:
        if field in updates:
            question[field] = updates[field]
    
    question['updated_at'] = datetime.utcnow().isoformat()
    
    # Save to persistent storage
    save_questions_db()
    
    logger.info("Question updated successfully", question_id=question_id)
    
    return APIResponse(
        success=True,
        message="Question updated",
        data=question
    )


@router.delete("/{question_id}", response_model=APIResponse)
async def delete_question(question_id: str):
    """Delete a question"""
    if question_id not in questions_db:
        raise HTTPException(status_code=404, detail="Question not found")
    
    del questions_db[question_id]
    
    # Save to persistent storage
    save_questions_db()
    
    return APIResponse(
        success=True,
        message="Question deleted"
    )


@router.post("/search", response_model=SearchResponse)
async def search_content(request: SearchRequest):
    """Search document content using semantic search"""
    rag_engine = get_rag_engine()
    
    results = await rag_engine.search(
        query=request.query,
        top_k=request.top_k,
        document_ids=request.document_ids,
        threshold=request.threshold
    )
    
    return SearchResponse(
        query=request.query,
        results=[
            SearchResult(
                chunk_id=r.chunk_id,
                document_id=r.document_id,
                content=r.content,
                score=r.score,
                metadata=r.metadata
            )
            for r in results
        ],
        total_results=len(results)
    )


@router.post("/export", response_model=APIResponse)
async def export_questions(request: ExportRequest):
    """Export questions in various formats"""
    # Determine export mode settings
    is_student_mode = request.export_mode == ExportMode.STUDENT
    
    # For student mode, always hide answers and explanations
    include_answers = False if is_student_mode else request.include_answers
    include_explanations = False if is_student_mode else request.include_explanations
    
    # Get questions
    questions = []
    for qid in request.question_ids:
        if qid in questions_db:
            q = questions_db[qid].copy()
            
            if not include_answers:
                q.pop('correct_answer', None)
                for opt in q.get('options', []):
                    opt.pop('is_correct', None)
            
            if not include_explanations:
                q.pop('explanation', None)
            
            questions.append(q)
    
    if not questions:
        raise HTTPException(status_code=404, detail="No questions found")
    
    # JSON export
    if request.format == ExportFormat.JSON:
        return APIResponse(
            success=True,
            data={
                'questions': questions,
                'format': 'json',
                'total': len(questions)
            }
        )
    
    # WORD/DOCX export
    if request.format in [ExportFormat.WORD, ExportFormat.DOCX]:
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            import base64
            from datetime import datetime
            
            doc = Document()
            
            if is_student_mode:
                # ========== STUDENT MODE - Exam Style ==========
                # Header - Exam title
                title = doc.add_heading('ĐỀ THI TRẮC NGHIỆM Y KHOA', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Exam info
                exam_info = doc.add_paragraph()
                exam_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                exam_info.add_run(f'Số câu hỏi: {len(questions)} câu')
                doc.add_paragraph()
                
                # Instructions
                instructions = doc.add_paragraph()
                instr_run = instructions.add_run('HƯỚNG DẪN: ')
                instr_run.bold = True
                instructions.add_run('Chọn một đáp án đúng nhất cho mỗi câu hỏi.')
                doc.add_paragraph()
                
                # Student info section
                info_para = doc.add_paragraph()
                info_para.add_run('Họ và tên: ').bold = True
                info_para.add_run('_' * 40)
                
                info_para2 = doc.add_paragraph()
                info_para2.add_run('Mã sinh viên: ').bold = True
                info_para2.add_run('_' * 30)
                info_para2.add_run('    ')
                info_para2.add_run('Lớp: ').bold = True
                info_para2.add_run('_' * 20)
                
                doc.add_paragraph()
                doc.add_paragraph('─' * 60)
                doc.add_paragraph()
                
                # Questions - clean exam format
                for i, q in enumerate(questions, 1):
                    # Question number with text
                    q_para = doc.add_paragraph()
                    q_num_run = q_para.add_run(f'Câu {i}. ')
                    q_num_run.bold = True
                    q_num_run.font.size = Pt(11)
                    q_para.add_run(q.get('question_text', ''))
                    
                    # Options - clean format without any marking
                    for opt in q.get('options', []):
                        opt_para = doc.add_paragraph()
                        opt_para.add_run(f"    {opt.get('id', '')}. {opt.get('text', '')}")
                    
                    doc.add_paragraph()  # Space between questions
                
                # Answer sheet section at the end
                doc.add_paragraph('─' * 60)
                answer_sheet_title = doc.add_paragraph()
                answer_sheet_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sheet_run = answer_sheet_title.add_run('BẢNG TRẢ LỜI')
                sheet_run.bold = True
                sheet_run.font.size = Pt(14)
                doc.add_paragraph()
                
                # Create answer grid (5 columns)
                table = doc.add_table(rows=(len(questions) + 4) // 5 + 1, cols=5)
                table.style = 'Table Grid'
                
                for i in range(len(questions)):
                    row_idx = i // 5
                    col_idx = i % 5
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = f'Câu {i+1}: ___'
                
                filename = f'de_thi_trac_nghiem_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
                
            else:
                # ========== TEACHER MODE - Full Export ==========
                # Title
                title = doc.add_heading('BỘ CÂU HỎI TRẮC NGHIỆM Y KHOA', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Subtitle with count
                subtitle = doc.add_paragraph(f'Tổng số: {len(questions)} câu hỏi')
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Empty line
                
                for i, q in enumerate(questions, 1):
                    # Question number and difficulty
                    difficulty_map = {'easy': 'Dễ', 'medium': 'Trung bình', 'hard': 'Khó'}
                    diff_text = difficulty_map.get(q.get('difficulty', 'medium'), 'Trung bình')
                    
                    q_header = doc.add_paragraph()
                    q_header_run = q_header.add_run(f'Câu {i}')
                    q_header_run.bold = True
                    q_header_run.font.size = Pt(12)
                    q_header.add_run(f' [{diff_text}]')
                    
                    # Topic if available
                    if q.get('topic'):
                        q_header.add_run(f' - {q.get("topic")}')
                    
                    # Question text (đã chứa scenario nếu là câu lâm sàng)
                    q_text = q.get('question_text', '')
                    q_para = doc.add_paragraph()
                    q_para.add_run(q_text)
                    
                    # Options
                    for opt in q.get('options', []):
                        opt_para = doc.add_paragraph()
                        opt_text = f"    {opt.get('id', '')}. {opt.get('text', '')}"
                        
                        if include_answers and opt.get('is_correct'):
                            opt_run = opt_para.add_run(opt_text + ' ✓')
                            opt_run.bold = True
                            opt_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                        else:
                            opt_para.add_run(opt_text)
                    
                    # Correct answer
                    if include_answers and q.get('correct_answer'):
                        answer_para = doc.add_paragraph()
                        answer_run = answer_para.add_run(f"➤ Đáp án đúng: {q.get('correct_answer')}")
                        answer_run.bold = True
                        answer_run.font.color.rgb = RGBColor(0, 100, 0)
                    
                    # Explanation
                    if include_explanations and q.get('explanation'):
                        exp_para = doc.add_paragraph()
                        exp_run = exp_para.add_run('💡 Giải thích: ')
                        exp_run.bold = True
                        exp_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                        exp_para.add_run(q.get('explanation', ''))
                    
                    # Add separator
                    doc.add_paragraph('─' * 50)
                
                filename = 'cau_hoi_trac_nghiem_giao_vien.docx'
            
            # Save to bytes
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            # Encode to base64
            doc_base64 = base64.b64encode(file_stream.read()).decode('utf-8')
            
            return APIResponse(
                success=True,
                data={
                    'file_content': doc_base64,
                    'filename': filename,
                    'format': 'word',
                    'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'total': len(questions),
                    'export_mode': 'student' if is_student_mode else 'teacher'
                }
            )
            
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="python-docx library not installed. Run: pip install python-docx"
            )
        except Exception as e:
            logger.error("Word export failed", error=str(e))
            raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")
    
    # TODO: Implement PDF, Excel export
    raise HTTPException(
        status_code=501,
        detail=f"Export format {request.format} not yet implemented"
    )


@router.get("/stats/overview", response_model=APIResponse)
async def get_question_stats():
    """Get question statistics"""
    questions = list(questions_db.values())
    
    # Count by difficulty
    by_difficulty = {}
    for q in questions:
        diff = q.get('difficulty', 'unknown')
        by_difficulty[diff] = by_difficulty.get(diff, 0) + 1
    
    # Count by type
    by_type = {}
    for q in questions:
        qtype = q.get('question_type', 'unknown')
        by_type[qtype] = by_type.get(qtype, 0) + 1
    
    # Count by topic
    by_topic = {}
    for q in questions:
        topic = q.get('topic', 'Unknown')
        by_topic[topic] = by_topic.get(topic, 0) + 1
    
    return APIResponse(
        success=True,
        data=QuestionStats(
            total_questions=len(questions),
            questions_by_difficulty=by_difficulty,
            questions_by_type=by_type,
            questions_by_topic=by_topic
        )
    )
